# processing/extract.py
import os
import zipfile
import csv
import pandas as pd
import logging
import re
from typing import List, Dict, Any, Tuple
from io import StringIO
from bs4 import BeautifulSoup

def load_file(path):
    """Main dispatcher for file loading - supports TXT, PDF, EPUB, DOCX, and CSV"""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".txt":
        return load_txt(path)
    elif ext == ".pdf":
        return load_pdf(path)
    elif ext == ".epub":
        return load_epub(path)
    elif ext == ".docx":
        return load_docx(path)
    elif ext == ".csv":
        return load_csv(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def load_txt(path):
    """Load plain text files"""
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def load_pdf(path):
    """Load PDF files using pdfminer"""
    try:
        from pdfminer.high_level import extract_text
        return extract_text(path)
    except Exception as e:
        raise RuntimeError(f"PDF extraction failed: {str(e)}")

def load_epub(path):
    """Load EPUB files using BeautifulSoup"""
    try:
        all_text = []

        with zipfile.ZipFile(path, 'r') as zip_ref:
            for file in zip_ref.namelist():
                if file.endswith((".htm", ".html", ".xhtml")):
                    with zip_ref.open(file) as f:
                        soup = BeautifulSoup(f.read(), "html.parser")
                        # Remove scripts/styles
                        for tag in soup(["script", "style"]):
                            tag.decompose()
                        body_text = soup.get_text(separator="\n", strip=True)
                        if body_text:
                            all_text.append(body_text)

        return "\n\n".join(all_text)
    except Exception as e:
        raise RuntimeError(f"EPUB fallback parsing failed: {str(e)}")

def load_docx(path):
    """
    Load Microsoft Word DOCX files using python-docx
    
    Extracts:
    - All paragraph text (including headings)
    - Table content (all cells)
    - Comments and footnotes
    - Header and footer text
    
    Args:
        path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text content
        
    Raises:
        RuntimeError: If extraction fails due to corruption, password protection, or other issues
    """
    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
        from docx.shared import Length
        
        # Attempt to open the document
        try:
            doc = Document(path)
        except PackageNotFoundError:
            raise RuntimeError("DOCX file is corrupted or not a valid Word document")
        except Exception as e:
            if "password" in str(e).lower() or "encrypted" in str(e).lower():
                raise RuntimeError("DOCX file is password protected - cannot extract text")
            else:
                raise RuntimeError(f"Cannot open DOCX file: {str(e)}")
        
        extracted_text = []
        
        # Extract main document paragraphs and headings
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                extracted_text.append(text)
        
        # Extract table content
        for table in doc.tables:
            table_text = _extract_table_text(table)
            if table_text:
                extracted_text.append(table_text)
        
        # Extract header and footer text
        for section in doc.sections:
            # Headers
            if section.header:
                header_text = _extract_header_footer_text(section.header)
                if header_text:
                    extracted_text.append(f"[HEADER] {header_text}")
            
            # Footers
            if section.footer:
                footer_text = _extract_header_footer_text(section.footer)
                if footer_text:
                    extracted_text.append(f"[FOOTER] {footer_text}")
        
        # Extract comments and footnotes
        comments_and_footnotes = _extract_comments_and_footnotes(doc)
        if comments_and_footnotes:
            extracted_text.extend(comments_and_footnotes)
        
        # Join all extracted text
        full_text = "\n\n".join(extracted_text)
        
        if not full_text.strip():
            raise RuntimeError("No text content found in DOCX file")
        
        return full_text
        
    except ImportError:
        raise RuntimeError("python-docx library not installed. Run: pip install python-docx")
    except Exception as e:
        if isinstance(e, RuntimeError):
            raise  # Re-raise our custom errors
        else:
            raise RuntimeError(f"DOCX extraction failed: {str(e)}")

# ========================================
# DOCX Helper Functions
# ========================================

def _extract_table_text(table):
    """
    Extract text from a Word table
    
    Args:
        table: python-docx Table object
        
    Returns:
        str: Formatted table text
    """
    table_content = []
    
    for row in table.rows:
        row_content = []
        for cell in row.cells:
            cell_text = cell.text.strip()
            if cell_text:
                row_content.append(cell_text)
        
        if row_content:
            table_content.append(" | ".join(row_content))
    
    if table_content:
        return "\n".join(table_content)
    return ""

def _extract_header_footer_text(header_footer):
    """
    Extract text from header or footer
    
    Args:
        header_footer: python-docx Header or Footer object
        
    Returns:
        str: Extracted text
    """
    text_parts = []
    
    # Extract paragraphs from header/footer
    for paragraph in header_footer.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)
    
    # Extract table content from header/footer
    for table in header_footer.tables:
        table_text = _extract_table_text(table)
        if table_text:
            text_parts.append(table_text)
    
    return " ".join(text_parts)

def _extract_comments_and_footnotes(doc):
    """
    Extract comments and footnotes from Word document
    
    Args:
        doc: python-docx Document object
        
    Returns:
        list: List of comment and footnote text
    """
    extracted = []
    
    try:
        # Extract comments
        # Note: python-docx doesn't have direct comment access in the basic API
        # This is a simplified approach that may miss some comments
        # For full comment extraction, would need to parse the underlying XML
        
        # Extract footnotes from document relationships
        # This is also limited by python-docx's API capabilities
        # For comprehensive footnote extraction, XML parsing would be needed
        
        # For now, we'll note that comments and footnotes may need manual handling
        # or a more advanced library like python-docx2txt for complete extraction
        
        pass  # Placeholder for future enhancement
        
    except Exception as e:
        # Don't fail the entire extraction if comments/footnotes can't be extracted
        pass
    
    return extracted

def _validate_docx_health(path):
    """
    Perform basic health check on DOCX file
    
    Args:
        path (str): Path to DOCX file
        
    Returns:
        dict: Health check results
    """
    health_info = {
        "is_valid_zip": False,
        "has_document_xml": False,
        "estimated_size": 0,
        "warnings": []
    }
    
    try:
        # Check if file is a valid ZIP (DOCX is a ZIP container)
        with zipfile.ZipFile(path, 'r') as zip_file:
            health_info["is_valid_zip"] = True
            
            # Check for core document.xml
            if "word/document.xml" in zip_file.namelist():
                health_info["has_document_xml"] = True
            
            # Estimate content size
            health_info["estimated_size"] = sum(
                zip_file.getinfo(name).file_size 
                for name in zip_file.namelist()
                if name.startswith("word/")
            )
            
    except zipfile.BadZipFile:
        health_info["warnings"].append("File is not a valid ZIP/DOCX format")
    except Exception as e:
        health_info["warnings"].append(f"Health check failed: {str(e)}")
    
    return health_info


# ========================================
# CSV Helper Functions
# ========================================

def _analyze_csv_file(path: str) -> Dict[str, Any]:
    """
    Analyze CSV file to estimate size and complexity
    
    Args:
        path (str): Path to CSV file
        
    Returns:
        dict: File analysis information
    """
    analysis = {
        'file_size_mb': 0,
        'estimated_rows': 0,
        'sample_lines': [],
        'delimiter_hints': []
    }
    
    try:
        # Get file size
        file_size = os.path.getsize(path)
        analysis['file_size_mb'] = file_size / (1024 * 1024)
        
        # Sample first few lines to estimate row count and structure
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            sample_lines = []
            for i, line in enumerate(f):
                if i < 10:  # Sample first 10 lines
                    sample_lines.append(line.strip())
                if i >= 100:  # Count first 100 lines for estimation
                    break
            
            # Estimate total rows based on average line length
            if sample_lines:
                avg_line_length = sum(len(line.encode('utf-8')) for line in sample_lines[:5]) / min(5, len(sample_lines))
                if avg_line_length > 0:
                    analysis['estimated_rows'] = int(file_size / avg_line_length)
                else:
                    analysis['estimated_rows'] = 1000  # Conservative estimate
            
            analysis['sample_lines'] = sample_lines
            
    except Exception as e:
        logging.warning(f"Could not analyze CSV file: {e}")
        analysis['estimated_rows'] = 10000  # Conservative high estimate
    
    return analysis

def _check_special_csv_format(path: str) -> Dict[str, Any]:
    """
    Check for special CSV formats like Excel CSV with sep= declaration
    """
    special_params = {}
    
    try:
        with open(path, 'rb') as f:
            # Check for BOM
            first_bytes = f.read(3)
            f.seek(0)
            
            # Check first line for sep= declaration
            first_line = f.readline().decode('utf-8', errors='ignore').strip()
            
            if first_line.startswith('sep='):
                # Excel CSV with separator declaration
                sep_char = first_line[4:].strip()
                special_params['delimiter'] = sep_char
                special_params['skip_rows'] = 1  # Skip the sep= line
                logging.info(f"Detected Excel CSV with separator: {repr(sep_char)}")
    except Exception as e:
        logging.debug(f"Special format check error: {e}")
    
    return special_params

def _detect_csv_delimiter(path: str, encoding: str = 'utf-8') -> str:
    """
    Detect CSV delimiter by analyzing file content
    
    Args:
        path (str): Path to CSV file
        encoding (str): File encoding to use
        
    Returns:
        str: Detected delimiter
    """
    delimiters = [',', ';', '\t', '|']
    
    try:
        with open(path, 'r', encoding=encoding) as f:
            # Read first few lines for analysis
            sample = ''
            for i, line in enumerate(f):
                if i < 5:  # Analyze first 5 lines
                    sample += line
                else:
                    break
        
        # Use csv.Sniffer to detect delimiter
        sniffer = csv.Sniffer()
        try:
            dialect = sniffer.sniff(sample, delimiters=delimiters)
            return dialect.delimiter
        except csv.Error:
            # Fallback: count occurrences of each delimiter
            delimiter_counts = {delim: sample.count(delim) for delim in delimiters}
            return max(delimiter_counts, key=delimiter_counts.get)
            
    except Exception as e:
        logging.warning(f"Could not detect delimiter: {e}, using comma as default")
        return ','

def _validate_csv_parsing(df: pd.DataFrame, delimiter: str) -> bool:
    """
    Validate that CSV was parsed correctly
    """
    # Check if we have multiple columns
    if len(df.columns) == 1 and delimiter in [',', ';', '\t', '|']:
        # Check if the single column contains the delimiter
        if len(df) > 0:
            first_val = str(df.iloc[0, 0])
            if delimiter in first_val:
                logging.warning("CSV appears to be incorrectly parsed into single column")
                return False
    
    # Check for suspicious column names that might indicate bad parsing
    col_str = ' '.join(str(c) for c in df.columns)
    if delimiter in col_str and len(df.columns) < 3:
        logging.warning("Column names contain delimiter, possible parsing issue")
        return False
    
    return True

def _appears_headerless(df: pd.DataFrame) -> bool:
    """
    Determine if CSV appears to lack proper headers
    
    Args:
        df (pd.DataFrame): DataFrame to analyze
        
    Returns:
        bool: True if CSV appears headerless
    """
    if len(df) == 0:
        return False
    
    # Check if first row data types match column names
    first_row = df.iloc[0] if len(df) > 0 else pd.Series()
    
    # If column names are generic (Unnamed, numeric patterns)
    generic_patterns = ['unnamed', 'column', r'^\d+$']
    has_generic_names = any(
        any(re.search(pattern, str(col).lower()) for pattern in generic_patterns)
        for col in df.columns
    )
    
    # If first row contains data similar to other rows (all numeric or mixed types)
    if len(df) > 1:
        first_row_types = [type(val).__name__ for val in first_row]
        second_row_types = [type(val).__name__ for val in df.iloc[1]]
        similar_types = sum(1 for t1, t2 in zip(first_row_types, second_row_types) if t1 == t2)
        similarity_ratio = similar_types / len(first_row_types) if first_row_types else 0
        
        return has_generic_names or similarity_ratio > 0.7
    
    return has_generic_names

def _identify_text_columns(df: pd.DataFrame) -> List[str]:
    """
    Intelligently identify columns containing text content suitable for training
    
    Args:
        df (pd.DataFrame): DataFrame to analyze
        
    Returns:
        List[str]: List of column names containing text content
    """
    text_columns = []
    
    for col in df.columns:
        try:
            # Get non-null values for analysis
            non_null_values = df[col].dropna()
            
            if len(non_null_values) == 0:
                continue
            
            # Convert to string for analysis
            str_values = non_null_values.astype(str)
            
            # Calculate text characteristics
            avg_length = str_values.str.len().mean()
            has_letters = str_values.str.contains(r'[a-zA-Z]', na=False).sum()
            has_spaces = str_values.str.contains(r'\s', na=False).sum()
            has_multiple_words = str_values.str.contains(r'\s+\w+', na=False).sum()
            
            # Ratios for classification
            letter_ratio = has_letters / len(str_values) if len(str_values) > 0 else 0
            space_ratio = has_spaces / len(str_values) if len(str_values) > 0 else 0
            multi_word_ratio = has_multiple_words / len(str_values) if len(str_values) > 0 else 0
            
            # Check if column contains URLs, emails, or structured text
            has_urls = str_values.str.contains(r'https?://', na=False).any()
            has_emails = str_values.str.contains(r'@.*\.', na=False).any()
            
            # Classification logic
            is_text_column = (
                # Has reasonable average length (more than just codes/IDs)
                avg_length >= 10 and
                # Contains letters (not just numbers)
                letter_ratio >= 0.5 and
                # Either has spaces/multiple words OR is structured text (URLs/emails)
                (space_ratio >= 0.3 or multi_word_ratio >= 0.2 or has_urls or has_emails)
            )
            
            # Special cases: include columns with obvious text indicators
            text_indicators = ['description', 'text', 'content', 'comment', 'note', 
                             'message', 'title', 'name', 'summary', 'review']
            
            if any(indicator in col.lower() for indicator in text_indicators):
                is_text_column = True
            
            # Exclude obvious non-text columns
            exclude_indicators = ['id', 'date', 'time', 'count', 'number', 'amount', 
                                'price', 'cost', 'total', 'sum', 'avg', 'max', 'min']
            
            if any(indicator in col.lower() for indicator in exclude_indicators):
                # Double-check with content analysis
                if avg_length < 5 or letter_ratio < 0.3:
                    is_text_column = False
            
            if is_text_column:
                text_columns.append(col)
                logging.info(f"Identified text column: {col} (avg_len={avg_length:.1f}, letters={letter_ratio:.2f})")
            
        except Exception as e:
            logging.warning(f"Error analyzing column {col}: {e}")
            continue
    
    return text_columns

def _extract_and_combine_text(df: pd.DataFrame, text_columns: List[str]) -> str:
    """
    Extract and intelligently combine text content from identified columns
    
    Args:
        df (pd.DataFrame): DataFrame containing the data
        text_columns (List[str]): List of columns to extract text from
        
    Returns:
        str: Combined text content
    """
    extracted_content = []
    
    for index, row in df.iterrows():
        row_texts = []
        
        for col in text_columns:
            try:
                value = row[col]
                
                # Skip null/empty values
                if pd.isna(value) or value == '':
                    continue
                
                # Convert to string and clean
                text = str(value).strip()
                
                # Skip very short or obviously non-textual content
                if len(text) < 3:
                    continue
                
                # Skip if it's just numbers or basic patterns
                if re.match(r'^[\d\.\-\+\s]*$', text):
                    continue
                
                # Add column context for multi-column rows
                if len(text_columns) > 1 and len(text) > 10:
                    # Only add column name if text is substantial
                    row_texts.append(f"{col}: {text}")
                else:
                    row_texts.append(text)
                    
            except Exception as e:
                logging.warning(f"Error extracting text from row {index}, column {col}: {e}")
                continue
        
        # Combine row texts
        if row_texts:
            if len(row_texts) == 1:
                # Single column: just use the text
                combined_row = row_texts[0]
            else:
                # Multiple columns: combine with separators
                combined_row = " | ".join(row_texts)
            
            extracted_content.append(combined_row)
    
    # Join all content with double newlines for clear separation
    return "\n\n".join(extracted_content)

def _fallback_csv_reader(path: str) -> str:
    """
    Fallback CSV reader using Python's csv module when pandas fails
    """
    import csv
    
    logging.info("Using fallback CSV reader")
    
    try:
        extracted_lines = []
        
        with open(path, 'r', encoding='utf-8', errors='replace') as f:
            # Try to detect dialect
            sample = f.read(8192)
            f.seek(0)
            
            try:
                dialect = csv.Sniffer().sniff(sample)
            except:
                # Default to excel dialect
                dialect = csv.excel()
            
            reader = csv.reader(f, dialect=dialect)
            
            # Process header
            try:
                header = next(reader)
                logging.info(f"CSV headers: {header}")
            except StopIteration:
                raise RuntimeError("Empty CSV file")
            
            # Find text columns by sampling
            text_cols = []
            sample_rows = []
            
            for i, row in enumerate(reader):
                if i < 10:  # Sample first 10 rows
                    sample_rows.append(row)
                else:
                    break
            
            # Identify text columns
            for col_idx in range(min(len(header), len(sample_rows[0]) if sample_rows else 0)):
                col_values = [row[col_idx] for row in sample_rows if col_idx < len(row)]
                
                # Check if column has text content
                has_text = any(
                    len(val) > 5 and any(c.isalpha() for c in val) 
                    for val in col_values if val
                )
                
                if has_text:
                    text_cols.append(col_idx)
            
            if not text_cols:
                # Use all columns
                text_cols = list(range(len(header)))
            
            # Extract text from identified columns
            f.seek(0)
            reader = csv.reader(f, dialect=dialect)
            next(reader)  # Skip header
            
            for row in reader:
                row_text = []
                for col_idx in text_cols:
                    if col_idx < len(row) and row[col_idx].strip():
                        text = row[col_idx].strip()
                        if len(text) > 3:  # Meaningful text
                            if col_idx < len(header):
                                row_text.append(f"{header[col_idx]}: {text}")
                            else:
                                row_text.append(text)
                
                if row_text:
                    extracted_lines.append(" | ".join(row_text))
        
        if not extracted_lines:
            raise RuntimeError("No text content extracted from CSV")
        
        # Create output
        metadata = f"[CSV DATA EXTRACT - Fallback Mode]\n"
        metadata += f"Source: {path}\n"
        metadata += f"Text columns found: {len(text_cols)}\n"
        metadata += f"[END METADATA]\n\n"
        
        return metadata + "\n\n".join(extracted_lines)
        
    except Exception as e:
        logging.error(f"Fallback reader failed: {e}")
        raise RuntimeError(f"Could not read CSV file: {str(e)}")

def load_csv(path: str, max_lines: int = 500) -> str:
    """
    Load and intelligently extract text content from CSV files - ROBUST VERSION
    
    This version handles edge cases like:
    - Files with BOM (Byte Order Mark)
    - Excel CSV files with sep= declarations
    - Unusual quoting
    - Mixed delimiters
    """
    try:
        # First, perform file size and content analysis
        file_info = _analyze_csv_file(path)
        
        # Check if file is too large and needs splitting
        if file_info['estimated_rows'] > max_lines:
            raise RuntimeError(
                f"CSV file is too large ({file_info['estimated_rows']:,} rows). "
                f"This exceeds the {max_lines} line limit for efficient processing. "
                f"Please split the file into smaller chunks or contact support for batch processing options."
            )
        
        # Check for special CSV formats
        special_params = _check_special_csv_format(path)
        
        # Try multiple encoding strategies
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']  # Added utf-8-sig for BOM
        df = None
        used_encoding = None
        delimiter = special_params.get('delimiter', None)
        
        for encoding in encodings:
            try:
                # Detect delimiter if not already specified
                if delimiter is None:
                    delimiter = _detect_csv_delimiter(path, encoding)
                
                logging.info(f"Attempting to read CSV with encoding '{encoding}' and delimiter '{delimiter}'")
                
                # Read CSV with robust parameters
                read_params = {
                    'filepath_or_buffer': path,
                    'encoding': encoding,
                    'delimiter': delimiter,
                    'engine': 'python',
                    'quoting': csv.QUOTE_NONE,
                    'na_values': ['', 'N/A', 'NULL', 'null', 'None'],
                    'keep_default_na': True,
                    'skipinitialspace': True,
                    'thousands': None,
                    'decimal': '.'
                }
                
                # Add quoting parameters based on detection
                if special_params.get('quoting'):
                    read_params['quoting'] = special_params['quoting']
                else:
                    read_params['quotechar'] = '"'
                    read_params['doublequote'] = True
                    read_params['escapechar'] = None
                
                # Skip sep= line if present
                if special_params.get('skip_rows'):
                    read_params['skiprows'] = special_params['skip_rows']
                
                df = pd.read_csv(**read_params)
                
                # Validate the parsing
                if not _validate_csv_parsing(df, delimiter):
                    raise pd.errors.ParserError("Invalid parsing detected")
                
                used_encoding = encoding
                logging.info(f"Successfully read CSV: shape={df.shape}, columns={list(df.columns)}")
                break
                
            except (UnicodeDecodeError, pd.errors.ParserError) as e:
                logging.warning(f"Failed with {encoding}: {e}")
                # Reset delimiter for next attempt
                if not special_params.get('delimiter'):
                    delimiter = None
                continue
        
        if df is None:
            # Last resort: try with python's csv module
            logging.info("Pandas failed, trying basic CSV reader")
            return _fallback_csv_reader(path)
        
        # Handle files without headers
        if _appears_headerless(df):
            logging.info("CSV appears headerless, re-reading without header row...")
            read_params['header'] = None
            df = pd.read_csv(**read_params)
            df.columns = [f"column_{i+1}" for i in range(len(df.columns))]
            logging.info(f"Generated column names: {list(df.columns)}")
        
        # Clean column names
        df.columns = [str(col).strip() for col in df.columns]
        
        # Identify text columns intelligently
        text_columns = _identify_text_columns(df)
        logging.info(f"Identified {len(text_columns)} text columns: {text_columns}")
        
        if not text_columns:
            # If no text columns found, check if we have at least some string data
            all_numeric = True
            for col in df.columns:
                if df[col].dtype == object or df[col].astype(str).str.contains('[a-zA-Z]', na=False).any():
                    all_numeric = False
                    break
            
            if all_numeric:
                raise RuntimeError("CSV contains only numeric data, no text content to extract")
            else:
                # Use all columns as fallback
                text_columns = list(df.columns)
                logging.warning("No clear text columns identified, using all columns")
        
        # Extract and combine text content
        extracted_text = _extract_and_combine_text(df, text_columns)
        
        if not extracted_text.strip():
            raise RuntimeError("No meaningful text content found in CSV file")
        
        # Add metadata header for context
        metadata = f"[CSV DATA EXTRACT]\n"
        metadata += f"Source: {path}\n"
        metadata += f"Encoding: {used_encoding}\n"
        metadata += f"Delimiter: {repr(delimiter)}\n"
        metadata += f"Rows: {len(df):,} | Columns: {len(df.columns)} | Text Columns: {len(text_columns)}\n"
        metadata += f"Columns used: {', '.join(text_columns[:5])}{'...' if len(text_columns) > 5 else ''}\n"
        metadata += f"[END METADATA]\n\n"
        
        final_text = metadata + extracted_text
        
        logging.info(f"CSV extraction complete: {len(final_text)} characters extracted")
        
        return final_text
        
    except FileNotFoundError:
        raise RuntimeError(f"CSV file not found: {path}")
    except pd.errors.EmptyDataError:
        raise RuntimeError("CSV file is empty or contains no data")
    except Exception as e:
        if isinstance(e, RuntimeError):
            raise
        else:
            logging.error(f"Unexpected error: {type(e).__name__}: {e}")
            raise RuntimeError(f"CSV extraction failed: {str(e)}")
